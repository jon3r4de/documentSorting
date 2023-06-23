from collections import deque
import os
import calendar
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import nsdecls
from docx.oxml import OxmlElement, parse_xml
from docx.shared import Pt, RGBColor
from docx import Document
import tkinter as tk
from tkinter import filedialog, messagebox
from tkinter import ttk
from PIL import Image, ImageTk

def merge_documents(parent_doc, sub_docs):
    for sub_doc in sub_docs:
        for element in sub_doc.element.body:
            parent_doc.element.body.append(element)


def parse_date(date_string):
    if len(date_string) < 8:
        return "Unknown"
    
    year = date_string[:4]
    month = date_string[4:6]
    day = date_string[6:8]

    if month == "00" and day == "00":
        if year == "0000":
            return "Undated"
        else:
            return year + " (Month Unknown)"
    elif day == "00":
        return "From " + calendar.month_name[int(month)] + " " + year
    else:
        return "From " + calendar.month_name[int(month)] + " " + day + ", " + year


def add_index(docu):
    table = docu.tables[0] # assume table is first and only table in the doc
    file_index = 1

    for row in table.rows[1:]:
        # Skip folder rows
        merged_cell = False
        for cell in row.cells:

            if cell.width != 0:
                merged_cell = True
                
        if merged_cell:
            continue

        row.cells[0].text = str(file_index)
        file_index += 1


def file_walk(root_dir, parent_doc=None, parent_table=None):
    if parent_doc is None:
        parent_doc = Document() 
        # parent_doc.add_paragraph('LIST OF DOCUMENTS', style='Heading 1').bold = True
        heading = parent_doc.add_paragraph('LIST OF DOCUMENTS')
        heading.style = 'Heading 1'
        run = heading.runs[0]
        run.bold = True
        run.font.size = Pt(28)
        run.font.name = 'Times New Roman'
        run.font.color.rgb = RGBColor(0, 0, 0)
        # run.font.color.rgb = parse_xml(r'<w:color {} w:val="000000"/>'.format(nsdecls('w')))
        heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER


        parent_table = parent_doc.add_table(rows=1, cols=3)
        parent_table.style = 'Table Grid'
        parent_table.allow_autofit = False
        parent_table.columns[0].width = int(Pt(1.0).twips)
        parent_table.columns[1].width = int(Pt(2.5).twips)
        parent_table.columns[2].width = int(Pt(4.0).twips)

        table_cells = parent_table.rows[0].cells
        table_cells[0].text = 'No.'
        table_cells[1].text = 'Date'
        table_cells[2].text = 'Description'

        for cell in table_cells:
            cell.width = Pt(240)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            cell.paragraphs[0].runs[0].font.size = Pt(14)
            cell.paragraphs[0].runs[0].font.bold = True

            # color_xml = """
            #     <w:r xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
            #         <w:rPr>
            #             <w:color w:val="000000"/>
            #         </w:rPr>
            #     </w:r>
            # """
            # run = cell.paragraphs[0].add_run()
            # run._r.append(parse_xml(color_xml))
            shading_elm = parse_xml(r'<w:shd {} w:fill="D9D9D9"/>'.format(nsdecls('w'))) # Dark grey background
            cell._tc.get_or_add_tcPr().append(shading_elm)

    items = os.listdir(root_dir)

    for item in items:
        # prevent macos stuff from printing
        if item == ".DS_Store":
            continue

        item_path = os.path.join(root_dir, item)

        if os.path.isfile(item_path):
            file_name = os.path.basename(item_path)
            file_name_parts = file_name.split(",")
            
            if len(file_name_parts) > 1:
                date_string = file_name_parts[0].strip()
                description = file_name_parts[1]

                row_cells = parent_table.add_row().cells
                row_cells[1].text = parse_date(date_string)
                row_cells[2].text = description

                for cell in row_cells:
                    shading_elm = parse_xml(r'<w:shd {} w:fill="FFFFFF"/>'.format(nsdecls('w')))
                    cell._tc.get_or_add_tcPr().append(shading_elm)

        else:
            folder_name = os.path.basename(item_path)
            
            row = parent_table.add_row().cells
            for cell in row:
                cell.width = Pt(240)
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

            start_cell = row[0].merge(row[2])
            start_cell.width = Pt(720)

            shading_elm = parse_xml(r'<w:shd {} w:fill="D9D9D9"/>'.format(nsdecls('w')))
            start_cell._tc.get_or_add_tcPr().append(shading_elm)

            start_cell.text = folder_name
            start_cell.paragraphs[0].alignment = WD_ALIGN_VERTICAL.CENTER

            sub_dir = os.path.join(root_dir, item)
            parent_doc = file_walk(sub_dir, parent_doc, parent_table)

    if parent_doc is not None and parent_table is not None:
        return parent_doc
    else:
        return Document()
 


def generate_file_index():
    root_directory = filedialog.askdirectory(title="Select Root Directory")

    if not root_directory:
        messagebox.showerror("Error", "No directory selected.")
        return

    output_path = filedialog.asksaveasfilename(title="Save Output File", defaultextension=".docx")

    if not output_path:
        messagebox.showerror("Error", "No output file selected.")
        return

    docu = file_walk(root_directory)
    add_index(docu)
    docu.save(output_path)
    messagebox.showinfo("Success", "File index generated successfully.")


def main():
    root = tk.Tk()
    root.title("File Index Generator")
    
        # Calculate the window position to center it on the screen
    window_width = 400
    window_height = 200
    screen_width = root.winfo_screenwidth()
    screen_height = root.winfo_screenheight()
    x = (screen_width - window_width) // 2
    y = (screen_height - window_height) // 4

    root.geometry(f"{window_width}x{window_height}+{x}+{y}")
    root.resizable(False, False)

    # Create a custom style for the button
    style = ttk.Style()
    style.configure("RoundedButton.TButton", font=("Arial", 12), borderwidth=0, relief="flat", background="#FFFFFF")
    style.map("RoundedButton.TButton", background=[("active", "#FFFFFF")])

    frame = ttk.Frame(root)
    frame.pack(pady=20)

    generate_button = ttk.Button(frame, text="Generate File Index", command=generate_file_index, style="RoundedButton.TButton")
    generate_button.pack()

    root.mainloop()


if __name__ == "__main__":
    main()
